from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
J = WD_ALIGN_PARAGRAPH.JUSTIFY; C = WD_ALIGN_PARAGRAPH.CENTER; L = WD_ALIGN_PARAGRAPH.LEFT

def build(doc):
    st = doc.styles['Normal']; st.font.name = 'Times New Roman'; st.font.size = Pt(12)
    for s in doc.sections:
        s.left_margin=Cm(2); s.right_margin=Cm(1.5); s.top_margin=Cm(1.5); s.bottom_margin=Cm(1.5)

def para(doc, runs=None, text=None, align=J, bold=False, size=12, after=6):
    p = doc.add_paragraph(); p.alignment = align; p.paragraph_format.space_after = Pt(after)
    if text is not None:
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    if runs:
        for t in runs:
            r = p.add_run(t); r.font.size = Pt(size)
    return p

def party(doc, pre, role, tail):
    # условный блок: организация/ИП ИЛИ физлицо
    org = ('{#%s_org}%s, ОГРН {%s_ogrn}, ИНН {%s_inn}, КПП {%s_kpp}, адрес: {%s_address}, '
           'в лице {%s_signer_role} {%s_signer_fio}, действующего на основании {%s_basis}, '
           'именуем{%s_end} в дальнейшем «%s»{/%s_org}') % (
           pre, '{%s_org_name}'%pre, pre,pre,pre,pre,pre,pre,pre,pre, role, pre)
    person = ('{^%s_org}Гражданин(ка) {%s_fio}, {%s_birth} г.р., место рождения: {%s_birthplace}, '
              'паспорт {%s_pasp_series} № {%s_pasp_number}, выдан {%s_pasp_issued_by} {%s_pasp_issued_date}, '
              'код подразделения {%s_pasp_code}, зарегистрированн{%s_end} по адресу: {%s_address}, '
              'именуем{%s_end} в дальнейшем «%s»{/%s_org}') % (
              pre,pre,pre,pre,pre,pre,pre,pre,pre,pre,pre,pre, role, pre)
    return para(doc, runs=[org, person, tail])

doc = Document(); build(doc)
para(doc, text='ДОГОВОР КУПЛИ-ПРОДАЖИ ТРАНСПОРТНОГО СРЕДСТВА', align=C, bold=True, size=14, after=2)
para(doc, runs=['г. {city}', '\t\t\t\t\t\t', '«{day}» {month} {year} г.'], align=L, after=12)

party(doc, 'seller', 'Продавец', ', с одной стороны, и')
party(doc, 'buyer', 'Покупатель', ', с другой стороны,')
para(doc, text='вместе именуемые «Стороны», заключили настоящий договор о нижеследующем:')

para(doc, text='1. ПРЕДМЕТ ДОГОВОРА', align=L, bold=True, after=4)
para(doc, text='1.1. Продавец обязуется передать в собственность Покупателя, а Покупатель — принять и оплатить '
     'транспортное средство (далее — ТС) со следующими характеристиками:')
specs = [('Марка, модель ТС','{car_brand}'),('Идентификационный номер (VIN)','{car_vin}'),
    ('Тип ТС','{car_type}'),('Категория ТС','{car_category}'),('Год выпуска','{car_year}'),
    ('Модель, № двигателя','{car_engine}'),('Шасси (рама) №','{car_chassis}'),
    ('Кузов (кабина, прицеп) №','{car_body}'),('Цвет кузова (кабины)','{car_color}'),
    ('Мощность двигателя, л.с. (кВт)','{car_power}'),('Рабочий объём двигателя, куб. см','{car_volume}'),
    ('Государственный регистрационный знак','{car_plate}'),
    ('Паспорт ТС (ПТС)','серия {pts_series} № {pts_number}, выдан {pts_issued}'),
    ('Свидетельство о регистрации (СТС)','серия {sts_series} № {sts_number}')]
t = doc.add_table(rows=0, cols=2); t.style='Table Grid'
for k,v in specs:
    c = t.add_row().cells; c[0].text=k; c[1].text=v
    c[0].paragraphs[0].runs[0].bold=True
    for cell in c:
        for pp in cell.paragraphs:
            pp.paragraph_format.space_after=Pt(0)
            for rr in pp.runs: rr.font.size=Pt(11)
para(doc, after=4)
para(doc, text='1.2. Продавец гарантирует, что до заключения настоящего договора ТС никому не продано, не заложено, '
     'в споре и под арестом (запрещением) не состоит, правами третьих лиц не обременено.')

para(doc, text='2. ЦЕНА ДОГОВОРА И ПОРЯДОК РАСЧЁТОВ', align=L, bold=True, after=4)
para(doc, text='2.1. Стоимость ТС составляет {price_num} ({price_words}) рублей.')
para(doc, text='2.2. Расчёт произведён полностью в момент подписания настоящего договора.')

para(doc, text='3. ПЕРЕДАЧА ТРАНСПОРТНОГО СРЕДСТВА', align=L, bold=True, after=4)
para(doc, text='3.1. Продавец передал, а Покупатель принял ТС, относящиеся к нему документы (ПТС, СТС) и ключи '
     'в момент подписания договора. Настоящий договор имеет силу акта приёма-передачи.')
para(doc, text='3.2. Право собственности на ТС переходит к Покупателю с момента подписания настоящего договора.')

para(doc, text='4. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ', align=L, bold=True, after=4)
para(doc, text='4.1. Договор составлен в трёх экземплярах равной юридической силы — по одному для каждой Стороны '
     'и один для органа ГИБДД.')
para(doc, text='4.2. Договор вступает в силу с момента подписания Сторонами.')

para(doc, text='5. ПОДПИСИ СТОРОН', align=L, bold=True, after=6)
sig = doc.add_table(rows=1, cols=2); sc = sig.rows[0].cells
def sigcell(cell, pre, role):
    cell.paragraphs[0].add_run(role+':').bold = True
    cell.add_paragraph('{#%s_org}{%s_org_name}, ИНН {%s_inn}{/%s_org}{^%s_org}{%s_fio}{/%s_org}' % (pre,pre,pre,pre,pre,pre,pre))
    cell.add_paragraph('{#%s_org}в лице {%s_signer_role} {%s_signer_fio}{/%s_org}' % (pre,pre,pre,pre))
    cell.add_paragraph()
    cell.add_paragraph('М.П. ___________ / {%s_short}' % pre)
    for pp in cell.paragraphs:
        pp.paragraph_format.space_after=Pt(2)
        for rr in pp.runs: rr.font.size=Pt(11)
sigcell(sc[0],'seller','ПРОДАВЕЦ'); sigcell(sc[1],'buyer','ПОКУПАТЕЛЬ')
para(doc, after=8)
para(doc, text='Деньги в сумме {price_num} рублей получил, ТС передал: ___________ / {seller_short}', align=L, size=11)
para(doc, text='ТС и документы получил, претензий не имею: ___________ / {buyer_short}', align=L, size=11)

doc.save('dkp-bot/template/dkp-template.docx')
print('template saved')
